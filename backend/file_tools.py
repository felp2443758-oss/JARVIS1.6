"""File conversion / extraction utilities.

Supported in MVP:
 - PDF -> text (pypdf)
 - DOCX -> text (python-docx)
 - Image -> text (OCR via Gemini vision)
 - Plain text passthrough

Returns metadata + extracted text.
"""
from __future__ import annotations
import io
import logging
from pathlib import Path
from typing import Optional, Dict, Any

from pypdf import PdfReader
from docx import Document

from vision_tools import analyze_image_bytes

logger = logging.getLogger("jarvis.files")


def extract_pdf_text(raw: bytes) -> str:
    reader = PdfReader(io.BytesIO(raw))
    pages = []
    for p in reader.pages:
        try:
            pages.append(p.extract_text() or "")
        except Exception:
            pages.append("")
    return "\n\n".join(pages).strip()


def extract_docx_text(raw: bytes) -> str:
    doc = Document(io.BytesIO(raw))
    parts = []
    for para in doc.paragraphs:
        parts.append(para.text)
    # tables
    for t in doc.tables:
        for row in t.rows:
            parts.append(" | ".join(cell.text for cell in row.cells))
    return "\n".join(parts).strip()


async def convert_file(filename: str, raw: bytes, mime_type: str = "", question: Optional[str] = None) -> Dict[str, Any]:
    """Extracts text from file. For images, does OCR via vision model."""
    name = (filename or "").lower()
    ext = Path(name).suffix.lower()
    out: Dict[str, Any] = {"filename": filename, "mime_type": mime_type, "ext": ext, "size": len(raw)}
    try:
        if ext == ".pdf" or "pdf" in (mime_type or ""):
            text = extract_pdf_text(raw)
            out["type"] = "pdf"
            out["text"] = text
            out["chars"] = len(text)
        elif ext in (".docx",) or "officedocument.wordprocessingml" in (mime_type or ""):
            text = extract_docx_text(raw)
            out["type"] = "docx"
            out["text"] = text
            out["chars"] = len(text)
        elif ext in (".txt", ".md", ".log", ".csv", ".json", ".xml", ".html", ".js", ".py", ".css") or (mime_type or "").startswith("text/"):
            text = raw.decode("utf-8", errors="ignore")
            out["type"] = "text"
            out["text"] = text
            out["chars"] = len(text)
        elif ext in (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp") or (mime_type or "").startswith("image/"):
            mt = mime_type or ("image/jpeg" if ext in (".jpg", ".jpeg") else f"image/{ext.lstrip('.')}")
            ocr = await analyze_image_bytes(raw, mime_type=mt, question=question or "Faça OCR completo do texto visível nesta imagem. Liste o texto fielmente, depois descreva brevemente o conteúdo visual.")
            out["type"] = "image"
            out["text"] = ocr
            out["chars"] = len(ocr)
        else:
            out["type"] = "unsupported"
            out["text"] = ""
            out["chars"] = 0
            out["error"] = f"Formato não suportado ({ext or mime_type})."
    except Exception as e:
        logger.exception("convert_file failed")
        out["type"] = out.get("type", "error")
        out["error"] = str(e)
        out["text"] = ""
        out["chars"] = 0
    return out
