"""J.A.R.V.I.S. Edge Agent — v2 action catalog (Windows-first, cross-platform where possible).

Each action returns a dict {ok, output|error}. Actions are dispatched from the
WebSocket command_handler on receipt of a `type=command` frame from the brain.

Security: `shell_exec` and `file_write` respect ALLOW_SHELL / ALLOW_FS envs (default off).
"""
from __future__ import annotations
import os
import sys
import base64
import shutil
import subprocess
import webbrowser
import platform
import socket
import tempfile
import time
from pathlib import Path
from typing import Any, Dict, Optional, List

IS_WIN = sys.platform.startswith("win")
IS_MAC = sys.platform == "darwin"

# Optional deps (import guarded — agent still runs on partial setup)
try:
    import pyautogui  # type: ignore
    HAS_PYAUTOGUI = True
except Exception:
    HAS_PYAUTOGUI = False

try:
    import psutil  # type: ignore
    HAS_PSUTIL = True
except Exception:
    HAS_PSUTIL = False

ALLOW_SHELL = os.getenv("JARVIS_ALLOW_SHELL", "0") == "1"
ALLOW_FS = os.getenv("JARVIS_ALLOW_FS", "1") == "1"

# -------------------- App launcher --------------------
# Windows: use `start` shell, common aliases resolve via Start Menu / PATH.
WIN_APP_ALIASES = {
    "spotify": "spotify",
    "chrome": "chrome",
    "edge": "msedge",
    "firefox": "firefox",
    "vscode": "code",
    "code": "code",
    "notepad": "notepad",
    "calc": "calc",
    "calculadora": "calc",
    "paint": "mspaint",
    "cmd": "cmd",
    "powershell": "powershell",
    "explorer": "explorer",
    "arquivos": "explorer",
    "word": "winword",
    "excel": "excel",
    "powerpoint": "powerpnt",
    "discord": "discord",
    "steam": "steam",
    "telegram": "telegram",
    "whatsapp": "whatsapp",
}


def _win_start(target: str) -> None:
    # cmd /c start "" TARGET opens whatever handler (URL, exe, ms-store URI) matches
    subprocess.Popen(["cmd", "/c", "start", "", target], shell=False)


def open_app(name: str) -> Dict[str, Any]:
    n = (name or "").strip().lower()
    if not n:
        return {"ok": False, "error": "nome do app vazio"}
    try:
        if IS_WIN:
            target = WIN_APP_ALIASES.get(n, n)
            _win_start(target)
        elif IS_MAC:
            subprocess.Popen(["open", "-a", name])
        else:
            subprocess.Popen([n])
        return {"ok": True, "output": f"Abrindo {name}"}
    except FileNotFoundError:
        # Fallback: try opening a web-app version for common services
        webmap = {
            "spotify": "https://open.spotify.com",
            "whatsapp": "https://web.whatsapp.com",
            "discord": "https://discord.com/app",
            "gmail": "https://mail.google.com",
        }
        if n in webmap:
            webbrowser.open(webmap[n])
            return {"ok": True, "output": f"App não instalado; abrindo versão web de {name}"}
        return {"ok": False, "error": f"App '{name}' não encontrado"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def close_app(name: str) -> Dict[str, Any]:
    if not HAS_PSUTIL:
        return {"ok": False, "error": "psutil não instalado"}
    n = (name or "").strip().lower()
    killed = 0
    for p in psutil.process_iter(["name", "pid"]):
        try:
            pname = (p.info.get("name") or "").lower()
            if n in pname:
                p.terminate()
                killed += 1
        except Exception:
            continue
    return {"ok": killed > 0, "output": f"{killed} processo(s) encerrado(s)"}


def list_apps() -> Dict[str, Any]:
    if not HAS_PSUTIL:
        return {"ok": False, "error": "psutil não instalado"}
    apps = []
    seen = set()
    for p in psutil.process_iter(["name", "pid"]):
        try:
            n = (p.info.get("name") or "").strip()
            if not n or n.lower() in seen:
                continue
            seen.add(n.lower())
            apps.append({"name": n, "pid": p.info.get("pid")})
        except Exception:
            continue
    apps.sort(key=lambda x: x["name"].lower())
    return {"ok": True, "output": apps[:200]}


# -------------------- URLs --------------------
def open_url(url: str, new_window: bool = False) -> Dict[str, Any]:
    if not url:
        return {"ok": False, "error": "url vazia"}
    try:
        if new_window:
            webbrowser.open_new(url)
        else:
            webbrowser.open(url)
        return {"ok": True, "output": f"Abrindo {url}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# -------------------- Keyboard / Mouse --------------------
def type_text(text: str, interval: float = 0.02) -> Dict[str, Any]:
    if not HAS_PYAUTOGUI:
        return {"ok": False, "error": "pyautogui não instalado"}
    pyautogui.typewrite(text, interval=interval)
    return {"ok": True, "output": f"{len(text)} chars digitados"}


def press_keys(keys: List[str]) -> Dict[str, Any]:
    if not HAS_PYAUTOGUI:
        return {"ok": False, "error": "pyautogui não instalado"}
    if not keys:
        return {"ok": False, "error": "lista de teclas vazia"}
    pyautogui.hotkey(*keys)
    return {"ok": True, "output": "+".join(keys)}


def hotkey(sequence: str) -> Dict[str, Any]:
    return press_keys([k.strip() for k in sequence.split("+") if k.strip()])


def screenshot() -> Dict[str, Any]:
    if not HAS_PYAUTOGUI:
        return {"ok": False, "error": "pyautogui não instalado"}
    img = pyautogui.screenshot()
    tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
    img.save(tmp.name, "PNG")
    with open(tmp.name, "rb") as f:
        data = f.read()
    try:
        os.unlink(tmp.name)
    except Exception:
        pass
    return {"ok": True, "output": {"b64": base64.b64encode(data).decode("ascii"), "mime": "image/png"}}


def mouse_click(x: int, y: int, button: str = "left") -> Dict[str, Any]:
    if not HAS_PYAUTOGUI:
        return {"ok": False, "error": "pyautogui não instalado"}
    pyautogui.click(x=x, y=y, button=button)
    return {"ok": True, "output": f"click {button} @ {x},{y}"}


def mouse_move(x: int, y: int, duration: float = 0.25) -> Dict[str, Any]:
    if not HAS_PYAUTOGUI:
        return {"ok": False, "error": "pyautogui não instalado"}
    pyautogui.moveTo(x, y, duration=duration)
    return {"ok": True, "output": f"move to {x},{y}"}


# -------------------- Volume --------------------
def volume(action: str = "up", steps: int = 3) -> Dict[str, Any]:
    if not HAS_PYAUTOGUI:
        return {"ok": False, "error": "pyautogui não instalado"}
    key = {"up": "volumeup", "down": "volumedown", "mute": "volumemute"}.get(action, "volumeup")
    for _ in range(max(1, min(steps, 20))):
        pyautogui.press(key)
    return {"ok": True, "output": f"vol {action} x{steps}"}


# -------------------- Files --------------------
def file_read(path: str, max_bytes: int = 256 * 1024) -> Dict[str, Any]:
    if not ALLOW_FS:
        return {"ok": False, "error": "acesso a arquivos desabilitado (JARVIS_ALLOW_FS=0)"}
    p = Path(path).expanduser()
    if not p.is_file():
        return {"ok": False, "error": f"arquivo não encontrado: {p}"}
    try:
        with open(p, "rb") as f:
            data = f.read(max_bytes)
        try:
            text = data.decode("utf-8")
            return {"ok": True, "output": {"path": str(p), "text": text, "bytes": len(data)}}
        except UnicodeDecodeError:
            return {"ok": True, "output": {"path": str(p), "b64": base64.b64encode(data).decode("ascii"), "bytes": len(data)}}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def file_write(path: str, content: str, mode: str = "overwrite") -> Dict[str, Any]:
    if not ALLOW_FS:
        return {"ok": False, "error": "escrita em arquivos desabilitada (JARVIS_ALLOW_FS=0)"}
    p = Path(path).expanduser()
    p.parent.mkdir(parents=True, exist_ok=True)
    flag = "w" if mode == "overwrite" else "a"
    try:
        with open(p, flag, encoding="utf-8") as f:
            f.write(content)
        return {"ok": True, "output": {"path": str(p), "bytes": len(content.encode("utf-8"))}}
    except Exception as e:
        return {"ok": False, "error": str(e)}


def doc_edit(path: str, instructions: str) -> Dict[str, Any]:
    """Very lightweight document editor. Supports .txt (append) and .docx (append paragraph).
    For full LLM-driven editing, brain sends `file_write` with the new content instead.
    """
    if not ALLOW_FS:
        return {"ok": False, "error": "escrita em arquivos desabilitada"}
    p = Path(path).expanduser()
    ext = p.suffix.lower()
    try:
        if ext == ".txt" or ext == "":
            with open(p, "a", encoding="utf-8") as f:
                f.write("\n" + instructions)
            return {"ok": True, "output": f"Adicionado {len(instructions)} chars em {p}"}
        if ext == ".docx":
            try:
                from docx import Document  # type: ignore
            except Exception:
                return {"ok": False, "error": "python-docx não instalado"}
            doc = Document(str(p)) if p.exists() else Document()
            doc.add_paragraph(instructions)
            doc.save(str(p))
            return {"ok": True, "output": f"Parágrafo adicionado em {p}"}
        return {"ok": False, "error": f"extensão não suportada: {ext}"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# -------------------- Shell --------------------
def shell_exec(command: str, timeout: float = 30.0) -> Dict[str, Any]:
    if not ALLOW_SHELL:
        return {"ok": False, "error": "shell desabilitado (JARVIS_ALLOW_SHELL=0)"}
    try:
        result = subprocess.run(
            command,
            shell=True,
            capture_output=True,
            text=True,
            timeout=timeout,
        )
        return {
            "ok": result.returncode == 0,
            "output": {
                "stdout": result.stdout[-8000:],
                "stderr": result.stderr[-2000:],
                "code": result.returncode,
            },
        }
    except subprocess.TimeoutExpired:
        return {"ok": False, "error": "timeout"}
    except Exception as e:
        return {"ok": False, "error": str(e)}


# -------------------- System info --------------------
def system_info() -> Dict[str, Any]:
    info: Dict[str, Any] = {
        "platform": platform.platform(),
        "os": platform.system(),
        "release": platform.release(),
        "hostname": socket.gethostname(),
        "python": sys.version.split()[0],
        "cwd": os.getcwd(),
    }
    if HAS_PSUTIL:
        vm = psutil.virtual_memory()
        info.update({
            "cpu_percent": psutil.cpu_percent(interval=0.2),
            "cpu_count": psutil.cpu_count(),
            "ram_total_gb": round(vm.total / (1024**3), 1),
            "ram_used_gb": round(vm.used / (1024**3), 1),
            "ram_percent": vm.percent,
        })
        try:
            bat = psutil.sensors_battery()
            if bat:
                info["battery_percent"] = bat.percent
                info["battery_plugged"] = bat.power_plugged
        except Exception:
            pass
    return {"ok": True, "output": info}
